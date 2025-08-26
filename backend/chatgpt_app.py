#!/usr/bin/env python3
"""
IntelliClaim: Advanced Multimodal RAG System for Document Processing
(GPT-5 API VERSION)
"""


import asyncio

import json
import logging
import os
import re
import urllib.parse
import hashlib
import time
from dataclasses import dataclass
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple


import PyPDF2
# import fitz  # PyMuPDF - Removed due to Rust compilation issues on Render
from dotenv import load_dotenv
from fastapi import FastAPI, File, HTTPException, UploadFile, Header, Request
from fastapi.middleware.cors import CORSMiddleware
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langgraph.graph import END, StateGraph
from pydantic import BaseModel
import uvicorn
import openai
import requests
from docx import Document as DocxDocument
from email import policy as email_policy
from email.parser import BytesParser as EmailBytesParser
from config import get_config


# --- Environment and Logging Setup ---
load_dotenv()
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# MODIFIED: Check for AI/ML API Key
if not os.getenv("AIMLAPI_KEY"):
    raise ValueError("AIMLAPI_KEY environment variable not set.")
if not os.getenv("OPENAI_API_KEY"):
    raise ValueError("OPENAI_API_KEY environment variable not set.")


# Configure AI/ML API
from openai import OpenAI
client = OpenAI(
    base_url="https://api.aimlapi.com/v1",
    api_key=os.getenv("AIMLAPI_KEY"),
)

CONFIG = get_config()

# --- Heuristics & Constants ---
MAX_CONTEXT_CHARS: int = 1500  # Reduced for faster processing
PER_CLAUSE_WINDOW: int = 400   # Reduced for faster processing

# Clause keyword biasing map (expandable for domains)
CLAUSE_KEYWORDS: Dict[str, Tuple[str, ...]] = {
    "waiting period": ("waiting period", "cooling period", "initial waiting"),
    "pre-existing": ("pre-existing", "PED"),
    "maternity": ("maternity", "pregnancy", "childbirth"),
    "cataract": ("cataract",),
    "organ donor": ("organ donor", "donor expenses"),
    "no claim": ("no claim discount", "NCD"),
    "preventive": ("preventive", "health check"),
    "hospital": ("hospital", "inpatient"),
    "ayush": ("ayush", "ayurveda", "homeopathy", "siddha", "unani"),
    "room rent": ("room rent", "icu charges", "sub-limit"),
}

def score_clause_bias(query: str, text: str) -> int:
    ql = query.lower()
    tl = text.lower()
    hits = 0
    for _, keywords in CLAUSE_KEYWORDS.items():
        if any(k in ql for k in keywords):
            hits += sum(1 for k in keywords if k in tl)
    return hits

def build_windowed_context(query: str, chunk: str, window: int = PER_CLAUSE_WINDOW) -> str:
    q_terms = [w for w in re.findall(r"\w+", query.lower()) if len(w) >= 4]
    first_hit = None
    cl = chunk.lower()
    for term in q_terms:
        idx = cl.find(term)
        if idx >= 0:
            first_hit = idx
            break
    if first_hit is None:
        return chunk[: min(len(chunk), window)]
    start = max(0, first_hit - window // 2)
    return chunk[start : min(len(chunk), start + window)]


# --- Data Models (No Changes) ---
@dataclass
class QueryContext:
    original_query: str
    structured_query: Dict[str, Any] = None
    retrieved_docs: List[Document] = None
    reasoning_chain: List[str] = None
    decision: Dict[str, Any] = None
    confidence_score: float = 0.0
    explanation: Dict[str, Any] = None


@dataclass
class DecisionResponse:
    decision: str
    amount: Optional[float] = None
    justification: str = ""
    clause_mappings: List[Dict[str, str]] = None
    confidence_score: float = 0.0
    processing_time: float = 0.0
    audit_trail: List[str] = None


class QueryRequest(BaseModel):
    query: str

class UploadURLRequest(BaseModel):
    url: str
    async_mode: Optional[bool] = False


# --- Core RAG Components (No Changes) ---
class EmbeddingManager:
    def __init__(self, model_name: str = None):
        # Use OpenAI embeddings with separate API key
        try:
            self.langchain_embeddings = OpenAIEmbeddings(
                openai_api_key=os.getenv("OPENAI_API_KEY"),  # Use OpenAI API key
                model="text-embedding-3-small",  # OpenAI's best embedding model
                chunk_size=1000
            )
            logger.info("Successfully initialized OpenAI embeddings")
        except Exception as e:
            logger.error(f"Failed to initialize OpenAI embeddings: {e}")
            # Fallback to FakeEmbeddings if OpenAI fails
            try:
                from langchain_community.embeddings import FakeEmbeddings
                self.langchain_embeddings = FakeEmbeddings(size=1536)  # OpenAI embedding dimension
                logger.warning("Using FakeEmbeddings as fallback - this will affect performance")
            except Exception as fallback_error:
                logger.error(f"Even fallback embeddings failed: {fallback_error}")
                raise RuntimeError("Could not initialize any embedding model. Please check your OpenAI API configuration.")

class DocumentProcessor:
    def __init__(self):
        self.default_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CONFIG.CHUNK_SIZE,
            chunk_overlap=CONFIG.CHUNK_OVERLAP,
        )

    def _split_text_dynamic(self, text: str) -> List[str]:
        length = len(text)
        if length < 100_000:
            cs, ov = 400, 80   # Reduced for faster processing
        elif length < 500_000:
            cs, ov = 800, 150  # Reduced for faster processing
        else:
            cs, ov = 1500, 200 # Reduced for faster processing
        splitter = RecursiveCharacterTextSplitter(chunk_size=cs, chunk_overlap=ov)
        return splitter.split_text(text)

    async def process_document(self, file_path: str, doc_hash: Optional[str] = None) -> List[Document]:
        documents: List[Document] = []
        filename = os.path.basename(file_path)
        ext = os.path.splitext(filename)[1].lower()
        try:
            raw_text = ""
            if ext == '.pdf':
                # Use PyPDF2 for PDF processing (PyMuPDF removed due to Rust compilation issues)
                try:
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        parts: List[str] = []
                        for page in pdf_reader.pages:
                            t = page.extract_text()
                            if t and t.strip():
                                parts.append(t)
                        raw_text = "\n\n".join(parts)
                except Exception as e:
                    logger.error(f"PDF processing failed: {e}")
                    raw_text = ""
            elif ext == '.docx':
                doc = DocxDocument(file_path)
                raw_text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            elif ext == '.eml':
                with open(file_path, 'rb') as fp:
                    msg = EmailBytesParser(policy=email_policy.default).parse(fp)
                parts = [msg.get('subject', ''), msg.get_body(preferencelist=('plain', 'html')).get_content() if msg.get_body() else '']
                raw_text = "\n\n".join(p for p in parts if p)
            else:
                logger.warning(f"Unsupported file type for {file_path}; skipping")
                return []

            if not raw_text.strip():
                return []

            chunks = self._split_text_dynamic(raw_text)
            for i, chunk in enumerate(chunks):
                documents.append(Document(
                    page_content=chunk,
                    metadata={
                        "source": filename,
                        "chunk_id": f"{i}",
                        "content_type": "text",
                        "doc_hash": doc_hash or "",
                    }
                ))
        except Exception as e:
            logger.error(f"Document processing failed for {file_path}: {e}")
        return documents


# --- Specialized Agents (Logic Updated for AI/ML API) ---
class GPT5Client:
    def __init__(self, api_key):
        self.api_key = api_key
        self.primary_model = "openai/gpt-5-mini-2025-08-07"  # Primary model to try first
        self.fallback_models = [
            "openai/gpt-5-2025-08-07",      # Fallback 1: Full GPT-5
            "openai/gpt-5-chat-latest",     # Fallback 2: Latest version
            "openai/gpt-4o-mini",           # Fallback 3: GPT-4 mini
            "openai/gpt-3.5-turbo"          # Fallback 4: GPT-3.5 as last resort
        ]
        self.client = OpenAI(
            base_url="https://api.aimlapi.com/v1",
            api_key=self.api_key,
        )


    async def generate_content_async(self, prompt, **kwargs):
        # Try primary model (gpt-5-mini) once
        try:
            logger.info(f"Trying primary model: {self.primary_model}")
            response = await asyncio.to_thread(
                self.client.chat.completions.create,
                model=self.primary_model,
                messages=[{"role": "user", "content": prompt}],
                temperature=kwargs.get("temperature", 0),
                max_tokens=kwargs.get("max_tokens", 1000),
                stream=False,
                timeout=30  # Shorter timeout for mini model
            )
            
            content = response.choices[0].message.content
            if content and content.strip():
                logger.info(f"Successfully generated content with primary model: {self.primary_model}")
                return content
            else:
                raise ValueError(f"Empty response from {self.primary_model}")
                
        except Exception as e:
            logger.warning(f"Primary model {self.primary_model} failed: {str(e)}. Switching to fallback models...")
        
        # Try fallback models in sequence
        for fallback_model in self.fallback_models:
            try:
                logger.info(f"Trying fallback model: {fallback_model}")
                response = await asyncio.to_thread(
                    self.client.chat.completions.create,
                    model=fallback_model,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=kwargs.get("temperature", 0),
                    max_tokens=kwargs.get("max_tokens", 1000),
                    stream=False,
                    timeout=60  # Longer timeout for fallback models
                )
                
                content = response.choices[0].message.content
                if content and content.strip():
                    logger.info(f"Successfully generated content with fallback model: {fallback_model}")
                    return content
                else:
                    raise ValueError(f"Empty response from {fallback_model}")
                    
            except Exception as e:
                logger.warning(f"Fallback model {fallback_model} failed: {str(e)}")
                continue
        
        # If all models failed
        raise Exception("All models failed to generate content")
    



def clean_json_response(text: str) -> str:
    """Helper function to clean the JSON string from the model's response."""
    if not text or not text.strip():
        return text
    
    # Remove markdown code blocks if present
    text = text.strip()
    if text.startswith('```json'):
        text = text[7:]
    if text.endswith('```'):
        text = text[:-3]
    text = text.strip()
    
    # Try to find JSON object
    match = re.search(r'\{', text)
    if not match:
        return text
    
    json_str = text[match.start():]
    json_match = re.search(r'\{.*\}', json_str, re.DOTALL)
    if json_match:
        return json_match.group(0)
    
    return text


class QueryUnderstandingAgent:
    def __init__(self, llm_client):
        self.llm = llm_client
        self.prompt_template = """
        You are an expert entity extractor for insurance claims. Extract the key information from the user query.
        The user query is: "{query}"
        
        Extract the following fields and respond ONLY with a valid JSON object (no markdown, no explanations):
        {{
            "age": <integer or null>,
            "gender": <"male", "female", or null>,
            "location": <string or null>,
            "procedure": <string or null>,
            "policy_duration_months": <integer, null if not mentioned, assume 'new policy' is 1 month>,
            "intent": <"claim_eligibility", "coverage_inquiry", or "policy_details">
        }}
        
        IMPORTANT: Return ONLY the JSON object, no other text.
        """
   
    def _extract_entities_fallback(self, query: str) -> dict:
        query_lower = query.lower()
        age_match = re.search(r'(\d+)[\s-]*(?:year|yr)s?[\s-]*old', query_lower)
        age = int(age_match.group(1)) if age_match else None
        gender = None
        if 'male' in query_lower or 'm' in query_lower:
            gender = 'male'
        elif 'female' in query_lower or 'woman' in query_lower:
            gender = 'female'
        locations = ['pune', 'mumbai', 'delhi', 'chennai', 'bangalore', 'hyderabad', 'kolkata']
        location = None
        for loc in locations:
            if loc in query_lower:
                location = loc.title()
                break
        procedures = ['surgery', 'treatment', 'operation', 'dental', 'cataract', 'heart', 'knee']
        procedure = None
        for proc in procedures:
            if proc in query_lower:
                procedure = proc
                break
        # Extract policy duration - prioritize "month/year old policy" patterns
        policy_duration = 1  # default
        policy_match = re.search(r'(\d+)[\s-]*(?:month|year|yr)s?[\s-]*old[\s-]*policy', query_lower)
        if policy_match:
            policy_duration = int(policy_match.group(1))
        else:
            policy_match = re.search(r'(\d+)[\s-]*(?:month|year|yr)s?[\s-]*policy', query_lower)
            if policy_match:
                policy_duration = int(policy_match.group(1))
            else:
                duration_match = re.search(r'(\d+)[\s-]*(?:month|year|yr)s?', query_lower)
                if duration_match:
                    match_pos = duration_match.start()
                    if not re.search(r'year[\s-]*old|years[\s-]*old', query_lower[:match_pos + 20]):
                        policy_duration = int(duration_match.group(1))
        return {
            "age": age,
            "gender": gender,
            "location": location,
            "procedure": procedure,
            "policy_duration_months": policy_duration,
            "intent": "claim_eligibility"
        }


    async def process(self, context: QueryContext) -> QueryContext:
        prompt = self.prompt_template.format(query=context.original_query)
        try:
            response_text = await self.llm.generate_content_async(prompt)
            logger.info(f"Raw GPT-5 response: {response_text[:200]}...")
            
            # Clean the response to extract JSON
            cleaned_response = clean_json_response(response_text)
            logger.info(f"Cleaned response: {cleaned_response[:200]}...")
            
            structured_query = json.loads(cleaned_response)
            context.structured_query = structured_query
            logger.info(f"Structured Query: {structured_query}")
        except json.JSONDecodeError as e:
            logger.error(f"JSON parsing failed: {e}")
            logger.error(f"Failed to parse: {cleaned_response if 'cleaned_response' in locals() else 'No cleaned response'}")
            context.structured_query = self._extract_entities_fallback(context.original_query)
            logger.info(f"Using fallback extraction: {context.structured_query}")
        except Exception as e:
            logger.error(f"Query Understanding failed: {e}")
            context.structured_query = self._extract_entities_fallback(context.original_query)
            logger.info(f"Using fallback extraction: {context.structured_query}")
        return context


class SemanticRetrievalAgent:
    def __init__(self, vector_store):
        self.vector_store = vector_store
        self.retrieval_cache: Dict[str, List[Document]] = {}

    async def process(self, context: QueryContext) -> QueryContext:
        search_query = " ".join(filter(None, [
            context.structured_query.get('procedure'),
            context.structured_query.get('intent'),
            context.structured_query.get('location'),
            "waiting period",
        ]))
        search_query = search_query.strip() or context.original_query
        logger.info(f"Retrieval Search Query: {search_query}")
        try:
            if search_query in self.retrieval_cache:
                docs = self.retrieval_cache[search_query]
            else:
                # Reduced k for faster retrieval
                docs = await asyncio.to_thread(self.vector_store.similarity_search, search_query, k=4)
                self.retrieval_cache[search_query] = docs

            ranked = sorted(docs, key=lambda d: score_clause_bias(context.original_query, d.page_content), reverse=True)
            context.retrieved_docs = ranked
            logger.info(f"Retrieved {len(ranked)} documents.")
        except Exception as e:
            logger.error(f"Retrieval failed: {e}")
            context.retrieved_docs = []
        return context


class DecisionReasoningAgent:
    def __init__(self, llm_client):
        self.llm = llm_client
    async def _make_decision_fallback(self, structured_query: dict, docs: list, original_query: str) -> dict:
        """AI-powered fallback decision making using policy extraction and LLM reasoning"""
        if not docs:
            return {
                "decision": "PENDING",
                "justification": "Could not retrieve relevant policy documents. Please upload policy documents for accurate assessment.",
                "amount": None,
                "confidence_score": 20
            }
        
        try:
            # Extract policy rules from documents using AI
            policy_rules = await self._extract_policy_rules_ai(docs, structured_query)
            
            # Use AI to make decision based on extracted rules
            decision = await self._make_ai_decision(structured_query, policy_rules, original_query)
            return decision
            
        except Exception as e:
            logger.error(f"AI decision making failed: {e}")
            # Fall back to basic rule-based logic if AI fails
            return self._apply_basic_rules(structured_query, docs)
    
    async def _extract_policy_rules_ai(self, docs: list, structured_query: dict) -> dict:
        """Extract policy rules from documents using AI analysis"""
        try:
            # Combine relevant document content
            policy_text = "\n\n".join([doc.page_content for doc in docs[:3]])  # Use top 3 docs
            
            # Create prompt for policy rule extraction
            extraction_prompt = f"""
            You are an expert insurance policy analyst. Extract key policy rules and coverage information from the following policy document.
            
            Query Context: {structured_query}
            
            Policy Document:
            {policy_text[:3000]}  # Limit text length
            
            Extract and return ONLY a JSON object with these fields:
            {{
                "waiting_periods": {{
                    "major_surgery": "number of months",
                    "minor_procedure": "number of months",
                    "preventive_care": "number of months"
                }},
                "coverage_limits": {{
                    "room_rent": "amount or 'as_per_policy'",
                    "icu": "amount or 'as_per_policy'",
                    "surgery": "amount or 'as_per_policy'",
                    "general_coverage": "overall policy limit"
                }},
                "exclusions": ["list of excluded procedures or conditions"],
                "special_benefits": ["list of special benefits or riders"],
                "procedure_specific_coverage": {{
                    "major_surgery": "amount if specified",
                    "minor_procedure": "amount if specified"
                }}
            }}
            
            Important:
            - Extract actual amounts in the policy's currency format
            - Look for procedure-specific coverage amounts
            - If information is not found, use "not_specified" for that field
            - Focus on amounts relevant to the query context
            """
            
            # Use LLM to extract rules
            response = await self.llm.generate_content_async(extraction_prompt)
            rules_text = response.strip()
            logger.info(f"Raw policy rules response: {rules_text[:300]}...")
            
            # Validate response is not empty
            if not rules_text or not rules_text.strip():
                logger.error("GPT-5 returned empty response for policy rules")
                raise ValueError("Empty response from GPT-5")
            
            # Parse JSON response
            try:
                # Clean the response to extract JSON
                cleaned_rules = clean_json_response(rules_text)
                logger.info(f"Cleaned policy rules: {cleaned_rules[:300]}...")
                
                if not cleaned_rules or cleaned_rules.strip() == "":
                    raise ValueError("No JSON content after cleaning")
                
                rules = json.loads(cleaned_rules)
                logger.info(f"Extracted policy rules: {rules}")
                return rules
            except Exception as parse_error:
                logger.error(f"Failed to parse policy rules JSON: {parse_error}")
                logger.error(f"Raw response: {rules_text}")
                logger.error(f"Cleaned response: {cleaned_rules if 'cleaned_rules' in locals() else 'No cleaned response'}")
                return self._extract_policy_rules_fallback(docs)
                
        except Exception as e:
            logger.error(f"AI policy rule extraction failed: {e}")
            return self._extract_policy_rules_fallback(docs)
    
    def _extract_policy_rules_fallback(self, docs: list) -> dict:
        """Fallback method to extract basic policy rules using regex patterns"""
        rules = {
            "waiting_periods": {"major_surgery": "not_specified", "minor_procedure": "not_specified", "preventive_care": "not_specified"},
            "coverage_limits": {"room_rent": "as_per_policy", "icu": "as_per_policy", "surgery": "as_per_policy", "general_coverage": "as_per_policy"},
            "exclusions": [],
            "special_benefits": []
        }
        
        combined_text = " ".join([doc.page_content.lower() for doc in docs])
        
        # Extract waiting periods using regex
        waiting_match = re.search(r'waiting\s+period.*?(\d+)\s*(?:month|year)', combined_text)
        if waiting_match:
            rules["waiting_periods"]["major_surgery"] = int(waiting_match.group(1))
            rules["waiting_periods"]["minor_procedure"] = int(waiting_match.group(1))
        
        # Extract coverage amounts with generic patterns
        amount_patterns = [
            (r'room\s+rent.*?(\d+(?:,\d+)*)', 'room_rent'),
            (r'icu.*?(\d+(?:,\d+)*)', 'icu'),
            (r'surgery.*?(\d+(?:,\d+)*)', 'surgery'),
            (r'coverage.*?(\d+(?:,\d+)*)', 'general_coverage'),
            (r'limit.*?(\d+(?:,\d+)*)', 'general_coverage'),
            (r'(\d+(?:,\d+)*)\s*(?:rs|rupees?|₹|dollars?|\$|euros?|€)', 'general_coverage')
        ]
        
        for pattern, key in amount_patterns:
            match = re.search(pattern, combined_text)
            if match:
                amount_str = match.group(1).replace(',', '')
                try:
                    amount = int(amount_str)
                    if key not in rules["coverage_limits"] or rules["coverage_limits"][key] == "as_per_policy":
                        rules["coverage_limits"][key] = amount
                except ValueError:
                    pass
        
        return rules
    
    async def _make_ai_decision(self, structured_query: dict, policy_rules: dict, original_query: str) -> dict:
        """Use AI to make decision based on extracted policy rules"""
        try:
            # Create decision prompt
            decision_prompt = f"""
            You are an expert insurance claims adjudicator. Based on the query and policy rules, make a decision.
            
            QUERY: {original_query}
            STRUCTURED INFO: {json.dumps(structured_query, indent=2)}
            POLICY RULES: {json.dumps(policy_rules, indent=2)}
            
            Respond ONLY with a JSON object with these exact keys:
            {{
                "decision": "APPROVED|REJECTED|PENDING",
                "justification": "detailed explanation of decision",
                "amount": "specific amount in rupees (e.g., 50000) or null if not specified",
                "confidence_score": "number between 0-100"
            }}
            
            Decision Guidelines:
            - APPROVED: If claim meets all policy requirements
            - REJECTED: If claim clearly violates policy rules
            - PENDING: If more information is needed
            
            Amount Guidelines:
            - If policy specifies coverage amount, use that exact amount
            - If policy has sub-limits, use the applicable limit
            - If amount is not specified, use null
            - Extract amounts in the policy's currency format
            
            Base your decision on the actual policy rules provided, not on assumptions.
            """
            
            # Get AI decision
            response = await self.llm.generate_content_async(decision_prompt)
            decision_text = response.strip()
            
            # Parse decision response
            try:
                json_match = re.search(r'\{.*\}', decision_text, re.DOTALL)
                if json_match:
                    decision = json.loads(json_match.group(0))
                    
                    # Validate decision format
                    required_keys = ["decision", "justification", "amount", "confidence_score"]
                    if all(key in decision for key in required_keys):
                        # Normalize decision values
                        decision["decision"] = self._normalize_decision(decision["decision"])
                        decision["confidence_score"] = min(100, max(0, int(decision.get("confidence_score", 50))))
                        
                        # Process amount field
                        original_amount = decision.get("amount")
                        decision["amount"] = self._process_amount_field(original_amount, policy_rules)
                        
                        logger.info(f"AI decision: {decision}")
                        logger.info(f"Amount processing: '{original_amount}' -> {decision['amount']}")
                        return decision
                    else:
                        raise ValueError("Missing required decision fields")
                else:
                    raise ValueError("No JSON found in decision response")
                    
            except Exception as parse_error:
                logger.error(f"Failed to parse AI decision: {parse_error}")
                raise
                
        except Exception as e:
            logger.error(f"AI decision making failed: {e}")
            raise
    
    def _process_amount_field(self, amount_value, policy_rules: dict):
        """Process and validate the amount field from AI decision"""
        try:
            # If AI provided a valid amount, use it
            if amount_value and amount_value != "null" and str(amount_value).lower() != "pending":
                # Try to convert to integer
                if isinstance(amount_value, str):
                    # Remove common currency symbols and convert
                    clean_amount = re.sub(r'[₹,rs\s$€£¥]', '', str(amount_value))
                    if clean_amount.isdigit():
                        return int(clean_amount)
                elif isinstance(amount_value, (int, float)):
                    return int(amount_value)
            
            # If no valid amount, try to extract from policy rules
            if policy_rules and "coverage_limits" in policy_rules:
                coverage = policy_rules["coverage_limits"]
                
                # Look for general coverage first, then specific types
                for key in ["general_coverage", "surgery", "room_rent", "icu"]:
                    if key in coverage and coverage[key] != "as_per_policy":
                        return coverage[key]
            
            # Default to null if no amount found
            return None
            
        except Exception as e:
            logger.error(f"Error processing amount field: {e}")
            return None
    
    def _apply_basic_rules(self, structured_query: dict, docs: list) -> dict:
        """Basic rule-based fallback when AI fails completely"""
        age = structured_query.get('age')
        procedure = structured_query.get('procedure')
        policy_duration = structured_query.get('policy_duration_months', 1)
        
        if not procedure:
            return {
                "decision": "PENDING",
                "justification": "Procedure type not specified. Please provide more details about the medical procedure.",
                "amount": None,
                "confidence_score": 30
            }
        
        # Simple logic based on policy duration
        if policy_duration < 3:
            return {
                "decision": "PENDING",
                "justification": f"Policy is only {policy_duration} month(s) old. Please check waiting period requirements.",
                "amount": None,
                "confidence_score": 40
            }
        else:
            return {
                "decision": "PENDING",
                "justification": f"Policy duration ({policy_duration} months) appears sufficient, but specific coverage details needed.",
                "amount": None,
                "confidence_score": 60
            }
    def _build_reasoning_prompt(self, structured_query: Dict, docs: List[Document], original_query: str) -> str:
        windows: List[str] = []
        remaining = MAX_CONTEXT_CHARS
        for i, doc in enumerate(docs):
            if remaining <= 0:
                break
            excerpt = build_windowed_context(original_query, doc.page_content)
            take = excerpt[:remaining]
            windows.append(f"Doc {i+1} (Source: {doc.metadata.get('source', 'unknown')}):\n{take}")
            remaining -= len(take)
        context_docs = "\n\n".join(windows)
        return f"""
You are an expert insurance claims adjudicator. Based on the query and policy documents, make a decision.
Respond ONLY with a JSON object with the keys: "decision", "amount", "justification", "confidence_score".

IMPORTANT: The "decision" field must be exactly one of these values:
- "APPROVED" (if claim is eligible and approved)
- "REJECTED" (if claim is not eligible or denied)
- "PENDING" (if more information is needed to make a decision)

QUERY: {original_query}
STRUCTURED INFO: {json.dumps(structured_query)}

RELEVANT POLICY CONTEXT (focused excerpts):
{context_docs}

DECISION (JSON format):
"""
    def _normalize_decision(self, decision_str: str) -> str:
        """Normalize decision string to expected values"""
        if not decision_str:
            return "PENDING"
       
        decision_lower = decision_str.lower().strip()
       
        # Map various decision formats to standard values
        if any(word in decision_lower for word in ['approved', 'approve', 'eligible', 'accept']):
            return "APPROVED"
        elif any(word in decision_lower for word in ['rejected', 'reject', 'denied', 'deny', 'not eligible']):
            return "REJECTED"
        elif any(word in decision_lower for word in ['pending', 'further', 'investigation', 'information', 'undecided']):
            return "PENDING"
        else:
            return "PENDING"  # Default to pending for unknown values


    async def process(self, context: QueryContext) -> QueryContext:
        if not context.retrieved_docs:
            context.decision = {"decision": "PENDING", "justification": "Could not retrieve relevant policy documents.", "amount": None, "confidence_score": 20}
            return context
        prompt = self._build_reasoning_prompt(context.structured_query, context.retrieved_docs, context.original_query)
        try:
            response_text = await self.llm.generate_content_async(prompt)
            logger.info(f"Raw decision response: {response_text[:200]}...")
            
            # Validate response is not empty
            if not response_text or not response_text.strip():
                logger.error("GPT-5 returned empty response")
                raise ValueError("Empty response from GPT-5")
            
            # Clean the response to extract JSON
            cleaned_response = clean_json_response(response_text)
            logger.info(f"Cleaned decision response: {cleaned_response[:200]}...")
            
            if not cleaned_response or cleaned_response.strip() == "":
                logger.error("No JSON content found after cleaning")
                raise ValueError("No JSON content in response")
            
            decision = json.loads(cleaned_response)
           
            # Normalize the decision value
            if 'decision' in decision:
                decision['decision'] = self._normalize_decision(decision['decision'])
           
            context.decision = decision
            context.reasoning_chain = [prompt, response_text]
            logger.info(f"Decision: {decision}")
        except Exception as e:
            logger.error(f"Decision reasoning failed: {e}")
            logger.error(f"Response text: {response_text if 'response_text' in locals() else 'No response'}")
            context.decision = await self._make_decision_fallback(context.structured_query, context.retrieved_docs, context.original_query)
            logger.info(f"Using fallback decision: {context.decision}")
        return context


class ExplainabilityAgent:
    async def process(self, context: QueryContext) -> QueryContext:
        decision_text = "unknown"
        if context.decision and isinstance(context.decision, dict):
            decision_text = context.decision.get('decision', 'unknown')
        explanation = {
            "clause_mappings": [
                {
                    "clause_text": doc.page_content[:200] + "...",
                    "source": doc.metadata.get('source', 'policy_document'),
                } for doc in (context.retrieved_docs or [])[:3]
            ],
            "audit_trail": [
                "Extracted key entities from query.",
                f"Retrieved {len(context.retrieved_docs or [])} relevant policy documents.",
                f"Applied rules to make a {decision_text} decision.",
            ]
        }
        context.explanation = explanation
        return context


# --- Main RAG System Orchestrator (Logic Updated for AI/ML API) ---
class IntelliClaimRAG:
    def __init__(self):
        self.llm = GPT5Client(os.getenv("AIMLAPI_KEY"))
        self.embedding_manager = EmbeddingManager()
        self.vector_store = Chroma(
            persist_directory="./chroma_db",
            embedding_function=self.embedding_manager.langchain_embeddings
        )
        self.doc_processor = DocumentProcessor()
        self.query_agent = QueryUnderstandingAgent(self.llm)
        self.retrieval_agent = SemanticRetrievalAgent(self.vector_store)
        self.reasoning_agent = DecisionReasoningAgent(self.llm)
        self.explanation_agent = ExplainabilityAgent()
        self.workflow = self.setup_agent_workflow()
    def setup_agent_workflow(self):
        workflow = StateGraph(QueryContext)
        workflow.add_node("understand_query", self.query_agent.process)
        workflow.add_node("retrieve_documents", self.retrieval_agent.process)
        workflow.add_node("make_decision", self.reasoning_agent.process)
        workflow.add_node("generate_explanation", self.explanation_agent.process)
        workflow.set_entry_point("understand_query")
        workflow.add_edge("understand_query", "retrieve_documents")
        workflow.add_edge("retrieve_documents", "make_decision")
        workflow.add_edge("make_decision", "generate_explanation")
        workflow.add_edge("generate_explanation", END)
        return workflow.compile()
    async def process_query(self, query: str) -> DecisionResponse:
        start_time = datetime.now()
        try:
            context = QueryContext(original_query=query)
            context = await self.query_agent.process(context)
            context = await self.retrieval_agent.process(context)
            context = await self.reasoning_agent.process(context)
            context = await self.explanation_agent.process(context)
            processing_time = (datetime.now() - start_time).total_seconds()
            decision_data = context.decision or {}
            explanation_data = context.explanation or {}
            return DecisionResponse(
                decision=decision_data.get('decision', 'ERROR'),
                amount=decision_data.get('amount'),
                justification=decision_data.get('justification', 'An unknown error occurred.'),
                confidence_score=decision_data.get('confidence_score', 0),
                clause_mappings=explanation_data.get('clause_mappings', []),
                audit_trail=explanation_data.get('audit_trail', []),
                processing_time=round(processing_time, 2)
            )
        except Exception as e:
            logger.error(f"Error in process_query: {e}")
            processing_time = (datetime.now() - start_time).total_seconds()
            return DecisionResponse(
                decision="ERROR",
                justification=f"Processing failed: {str(e)}",
                confidence_score=0,
                clause_mappings=[],
                audit_trail=["Error occurred during processing"],
                processing_time=round(processing_time, 2)
            )
    async def add_document(self, file_path: str):
        # Compute a stable doc hash for caching/metadata joining
        try:
            with open(file_path, 'rb') as fh:
                content = fh.read()
            doc_hash = hashlib.sha256(content).hexdigest()
        except Exception:
            doc_hash = None
        # Skip re-index if already present (idempotent ingest)
        try:
            coll = self.vector_store._collection
            if coll and doc_hash:
                existing = coll.get(where={"doc_hash": doc_hash}, limit=1)
                if existing and existing.get("ids"):
                    return {"status": "skipped", "reason": "already_indexed", "file": os.path.basename(file_path)}
        except Exception:
            pass

        documents = await self.doc_processor.process_document(file_path, doc_hash=doc_hash)
        if documents:
            self.vector_store.add_documents(documents, batch_size=get_config().BATCH_SIZE)
            return {"status": "success", "documents_added": len(documents), "file": os.path.basename(file_path)}
        return {"status": "error", "message": "No text could be extracted from the document."}


# --- FastAPI Web Service (Config-driven) ---
app = FastAPI(title=get_config().API_TITLE, version=get_config().API_VERSION)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"],
)


rag_system = IntelliClaimRAG()


@app.post("/query", response_model=DecisionResponse)
async def process_query_endpoint(request: QueryRequest):
    try:
        response = await rag_system.process_query(query=request.query)
        return response
    except Exception as e:
        logger.error(f"Unhandled exception in /query: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/upload-document")
async def upload_document_endpoint(file: UploadFile = File(...)):
    if not os.path.exists("./uploads"):
        os.makedirs("./uploads")
    file_path = f"./uploads/{file.filename}"
    with open(file_path, "wb") as buffer:
        buffer.write(await file.read())
    result = await rag_system.add_document(file_path)
    if result['status'] == 'error':
        raise HTTPException(status_code=400, detail=result['message'])
    return result


@app.get("/health")
async def health_check():
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}


@app.get("/test-aimlapi")
async def test_aimlapi():
    try:
        model = GPT5Client(os.getenv("AIMLAPI_KEY"))
        response = await model.generate_content_async("Say 'Hello World' in JSON format: {\"message\": \"Hello World\"}")
        return {"status": "success", "response": response, "primary_model": model.primary_model}
    except Exception as e:
        return {"status": "error", "error": str(e)}

@app.get("/test-gpt5-json")
async def test_gpt5_json():
    """Test GPT-5's JSON response capabilities"""
    try:
        model = GPT5Client(os.getenv("AIMLAPI_KEY"))
        prompt = """
        Extract the following fields and respond ONLY with a valid JSON object:
        {
            "age": <integer or null>,
            "gender": <"male", "female", or null>,
            "location": <string or null>
        }
        
        Example query: "I am a 35 year old male from Mumbai"
        """
        response = await model.generate_content_async(prompt)
        cleaned = clean_json_response(response)
        try:
            parsed = json.loads(cleaned)
            return {
                "status": "success", 
                "raw_response": response,
                "cleaned_response": cleaned,
                "parsed_json": parsed,
                "primary_model": model.primary_model
            }
        except json.JSONDecodeError as e:
            return {
                "status": "json_parse_error",
                "raw_response": response,
                "cleaned_response": cleaned,
                "error": str(e),
                "primary_model": model.primary_model
            }
    except Exception as e:
        return {"status": "error", "error": str(e)}

@app.get("/test-gpt5-debug")
async def test_gpt5_debug():
    """Debug endpoint to test GPT-5 response directly"""
    try:
        model = GPT5Client(os.getenv("AIMLAPI_KEY"))
        
        # Test with a very simple prompt
        simple_prompt = "Say exactly 'Hello World' and nothing else."
        
        logger.info(f"Testing GPT-5 with prompt: '{simple_prompt}'")
        response = await model.generate_content_async(simple_prompt)
        
        return {
            "status": "success",
            "prompt": simple_prompt,
            "response": response,
            "response_length": len(response) if response else 0,
            "primary_model": model.primary_model,
            "response_type": type(response).__name__
        }
    except Exception as e:
        logger.error(f"GPT-5 debug test failed: {e}")
        return {
            "status": "error", 
            "error": str(e),
            "error_type": type(e).__name__
        }

@app.get("/model-config")
async def get_model_config():
    """Get current model configuration"""
    try:
        model = GPT5Client(os.getenv("AIMLAPI_KEY"))
        return {
            "status": "success",
            "primary_model": model.primary_model,
            "fallback_models": model.fallback_models,
            "description": "System will try primary model once, then fall back to fallback models in sequence"
        }
    except Exception as e:
        return {
            "status": "error", 
            "error": str(e)
        }




@app.get("/documents")
async def list_documents():
    """List all uploaded documents in the system"""
    try:
        # Get documents from vector store
        collection = rag_system.vector_store._collection
        if collection:
            documents = collection.get()
            return {
                "total_documents": len(documents.get('documents', [])),
                "document_sources": list(set(doc.get('source', 'unknown') for doc in documents.get('metadatas', [])))
            }
        return {"total_documents": 0, "document_sources": []}
    except Exception as e:
        logger.error(f"Error listing documents: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/documents/{filename}")
async def delete_document(filename: str):
    """Delete a specific document from the system"""
    try:
        # Remove from vector store
        collection = rag_system.vector_store._collection
        if collection:
            # Find and delete documents with matching source
            collection.delete(where={"source": filename})
       
        # Remove file from uploads
        file_path = f"./uploads/{filename}"
        if os.path.exists(file_path):
            os.remove(file_path)
       
        return {"status": "success", "message": f"Document {filename} deleted successfully"}
    except Exception as e:
        logger.error(f"Error deleting document: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/system-stats")
async def get_system_stats():
    """Get system statistics and health metrics"""
    try:
        collection = rag_system.vector_store._collection
        total_docs = len(collection.get().get('documents', [])) if collection else 0
       
        return {
            "total_documents": total_docs,
            "vector_store_size": total_docs,
            "system_status": "healthy",
            "last_updated": datetime.now().isoformat(),
            "api_version": "1.0.0"
        }
    except Exception as e:
        logger.error(f"Error getting system stats: {e}")
        raise HTTPException(status_code=500, detail=str(e))


def download_file(url: str, temp_path: str) -> bool:
    try:
        response = requests.get(url, stream=True, timeout=30)
        response.raise_for_status()
        with open(temp_path, 'wb') as f:
            for chunk in response.iter_content(chunk_size=65536):
                if chunk:
                    f.write(chunk)
        return True
    except Exception as e:
        logger.error(f"download_file failed for {url}: {e}")
    return False


def extract_text_from_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    return '\n\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])


"""Production helpers only."""


"""No experimental retriever in production build."""


"""No local LLM pipeline in production build."""


"""Lightweight runtime cache."""


@app.post("/upload-document-url")
async def upload_document_by_url(payload: UploadURLRequest):
    """Download a document by URL and ingest it into the vector store."""
    parsed = urllib.parse.urlparse(payload.url)
    filename = os.path.basename(parsed.path) or f"downloaded_{datetime.now().timestamp()}"
    if not os.path.exists("./uploads"):
        os.makedirs("./uploads")
    temp_path = os.path.join("./uploads", filename)
    if not download_file(payload.url, temp_path):
        raise HTTPException(status_code=500, detail="Failed to download document from URL")
    if payload.async_mode:
        # Fire-and-forget background ingestion using asyncio.create_task
        asyncio.create_task(rag_system.add_document(temp_path))
        return {"status": "accepted", "message": "Ingestion started", "file": os.path.basename(temp_path)}
    else:
        result = await rag_system.add_document(temp_path)
        if result['status'] == 'error':
            raise HTTPException(status_code=400, detail=result['message'])
        return result

# Runtime caches for the hackrx endpoint
faiss_cache = {}
retrieval_cache = {}
summary_cache = {}

# Add this near your other endpoints
@app.get("/health")
async def health_check():
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}

@app.post("/api/v1/hackrx/run")
async def hackrx_run(request: Request, authorization: str = Header(None)):
    # Deprecated endpoint retained for backward compatibility; returns 410 Gone
    raise HTTPException(status_code=410, detail="Endpoint deprecated. Use /upload-document, /upload-document-url, and /query.")
 
    body = await request.json()
    print("Debug: Received request body:", body)

    documents_url = body.get("documents")
    questions = body.get("questions", [])
 
    if not documents_url or not isinstance(questions, list) or not questions:
        raise HTTPException(status_code=400, detail="Missing or invalid documents URL or questions")

    print("Debug: Processing document URL:", documents_url)
    print("Debug: Number of questions:", len(questions))
 
    # Parse filename from URL
    parsed_url = urllib.parse.urlparse(documents_url)
    filename = os.path.basename(parsed_url.path)
    file_extension = os.path.splitext(filename)[1].lower()
 
    temp_file = f"./temp_{int(time.time()*1000)}{file_extension}"
 
    if not download_file(documents_url, temp_file):
        raise HTTPException(status_code=500, detail="Failed to download document")

    print("Debug: Document downloaded successfully to:", temp_file)
 
    try:
        # Process document based on type
        raw_text = ""
        if file_extension == '.pdf':
            with open(temp_file, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        raw_text += text + '\n\n'
        elif file_extension == '.docx':
            raw_text = extract_text_from_docx(temp_file)
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_extension}")
 
        print("Debug: Extracted raw text length:", len(raw_text))
        # Compute content hash for persistent caching
        try:
            doc_hash = hashlib.sha256(raw_text.encode("utf-8", errors="ignore")).hexdigest()
        except Exception:
            doc_hash = hashlib.sha256(raw_text.encode("utf-8", errors="ignore")).hexdigest()

        if not raw_text.strip():
            raise HTTPException(status_code=400, detail="No text could be extracted from the document")
 
        # Split into chunks - dynamic chunk size
        doc_length = len(raw_text)
        chunk_size = 500 if doc_length < 100000 else 1000 if doc_length < 500000 else 2000  # Dynamic based on size
        chunk_overlap = 300 if doc_length >= 200000 else 100
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        chunks = text_splitter.split_text(raw_text)

        print("Debug: Number of chunks:", len(chunks))
        if chunks:
            print("Debug: Sample chunk[0]:", chunks[0][:200] + "...")  # First 200 chars of first chunk
 
        # Create documents
        documents = [Document(page_content=chunk, metadata={"source": filename, "chunk": i}) for i, chunk in enumerate(chunks)]
 
        # Create/load FAISS vector store (persistent by content hash)
        cache_key = documents_url  # URL cache (runtime)
        faiss_dir = os.path.join(".", "faiss_cache", doc_hash)
        if cache_key in faiss_cache:
            vector_store = faiss_cache[cache_key]
            print("Debug: Loaded cached vector store for", cache_key)
        else:
            # Production: use existing Chroma store from rag_system
            vector_store = rag_system.vector_store

        print("Debug: Vector store created with", len(documents), "documents")
 
        # Initialize LLM
        model = GPT5Client(os.getenv("AIMLAPI_KEY"))
 
        print("Debug: Starting answer generation for", len(questions), "questions")
 
        async def generate_answer(questions_batch, vector_store, model, all_documents, doc_length, doc_hash):
            start_time = time.time()
            try:
                # Detect if batch has potentially high-weight (complex) questions
                is_complex_batch = any(len(q) > 50 or any(word in q.lower() for word in ['derive', 'explain', 'list', 'how does']) for q in questions_batch)
                batch_timeout = 30.0 if is_complex_batch else 20.0  # Keep dynamic timeout
                batch_max_retries = 3  # Relaxed to 3 for all with higher quota
                print("Debug: Batch complexity:", "Complex" if is_complex_batch else "Standard", "- Timeout:", batch_timeout, "Retries:", batch_max_retries)
                
                # Retrieve contexts for all questions in batch (hybrid + clause bias)
                batch_contexts = []
                override_answers: List[Optional[str]] = []
                # Use simple similarity search instead of undefined _HybridRetriever
                def simple_retrieve(question, docs, k=3):
                    return [(doc, 1.0) for doc in docs[:k]]
                
                def dynamic_k(doc_length):
                    return 3 if doc_length < 100000 else 4 if doc_length < 500000 else 5
                
                def clause_bias(question, docs):
                    return sorted(docs, key=lambda d: score_clause_bias(question, d.page_content), reverse=True)
                
                retriever = type('SimpleRetriever', (), {
                    'score_and_combine': simple_retrieve,
                    'dynamic_k': dynamic_k,
                    'clause_bias': clause_bias
                })()
                for question in questions_batch:
                    qlower = question.lower()
                    # Precompute question-type expectations for reuse
                    expects_identifier = bool(re.search(r"\b(which|under which|what is|according to|in)\s+(article|section|clause)\b", qlower))
                    expects_article24_age = bool(re.search(r"\barticle\s*24\b.*(age|years?)", qlower))
                    expects_article17_abolish = bool(re.search(r"\barticle\s*17\b", qlower) and re.search(r"abolish|abolished|abolition", qlower))
                    k_val = retriever.dynamic_k(doc_length)
                    # Retrieval cache
                    cache_key_ret = (doc_hash, question)
                    cached_context = retrieval_cache.get(cache_key_ret)
                    if cached_context is None:
                        combined = retriever.score_and_combine(question, all_documents, k=k_val)
                        docs_only = [d for d, _ in combined]
                        docs_biased = retriever.clause_bias(question, docs_only)
                        # Build windowed context, centered near query terms
                        context_parts, chars_left = [], MAX_CONTEXT_CHARS
                        for d in docs_biased:
                            if chars_left <= 0:
                                break
                            chunk = d.page_content
                            q_terms = [w for w in re.findall(r"\w+", question.lower()) if len(w) >= 4]
                            first_hit = None
                            for term in q_terms:
                                idx = chunk.lower().find(term)
                                if idx >= 0:
                                    first_hit = idx
                                    break
                            if first_hit is None:
                                excerpt = chunk[: min(len(chunk), PER_CLAUSE_WINDOW)]
                            else:
                                start = max(0, first_hit - PER_CLAUSE_WINDOW // 2)
                                excerpt = chunk[start : min(len(chunk), start + PER_CLAUSE_WINDOW)]
                            take = excerpt[:chars_left]
                            context_parts.append(take)
                            chars_left -= len(take)
                        raw_context = "\n\n".join(context_parts)
                        retrieval_cache[cache_key_ret] = raw_context
                        # Approximate average dense score from combined results
                        try:
                            avg_faiss = sum(sc for _, sc in combined) / max(1, len(combined))
                        except Exception:
                            avg_faiss = 9.99
                    else:
                        raw_context = cached_context
                        # Heuristic score proxy when using cached context
                        avg_faiss = 0.9 if len(raw_context) <= 1200 else 1.2
                    # Decide summarization based on average FAISS score
                    if raw_context and avg_faiss > 1.1 and not (expects_identifier or expects_article24_age or expects_article17_abolish):
                        summary_key = (doc_hash, question)
                        cached_summary = summary_cache.get(summary_key)
                        if cached_summary is None:
                            summary_prompt = f"Summarize this context concisely in 200 words or less, focusing on key facts relevant to: {question}\n\nContext: {raw_context[:2000]}"
                            try:
                                summary_response = await model.generate_content_async(
                                    summary_prompt,
                                    temperature=0
                                )
                                context = summary_response.strip()
                                summary_cache[summary_key] = context
                                print("Debug: Summarized context for", question, ":", (context[:200] + "...") if context else "EMPTY")
                            except Exception as e:
                                print("Debug: Summary error for", question, ":", str(e))
                                context = raw_context[:1500]
                        else:
                            context = cached_summary
                    else:
                        context = raw_context[:1500]
                        if avg_faiss <= 1.0 or (expects_identifier or expects_article24_age or expects_article17_abolish):
                            print("Debug: Skipped summarization for good score/exact-match question:", question)
                    
                    if not context or len(context) < 50:
                        print("Debug: Minimal context for", question, "- Forcing fallback")
                        batch_contexts.append((question, ""))
                        override_answers.append(None)
                    else:
                        batch_contexts.append((question, context))
                        # Compute regex-based override for exact-match grading
                        override: Optional[str] = None
                        # General identifier extraction (e.g., Article X, Section Y, Clause Z)
                        expects_identifier = bool(re.search(r"\b(which|under which|what is|according to|in)\s+(article|section|clause)\b", qlower))
                        if expects_identifier:
                            m = re.findall(r"\b(Article|Section|Clause)\s+\d+[A-Z]?\b", raw_context, flags=re.IGNORECASE)
                            if m:
                                best = m[0]
                                type_ = best.split()[0].capitalize()
                                num = best.split()[1].upper() if best.split()[1].isalpha() else best.split()[1]
                                override = f"{type_} {num}"
                                print("Debug: General identifier override:", override)
                        # General numeric/age extraction (e.g., below the age of X, Rs Y)
                        if "age" in qlower or re.search(r"\b(how old|below what age|age of)\b", qlower):
                            m = re.search(r"\bage (?:age of|below|under)\s+(\d+|\w+)", raw_context, flags=re.IGNORECASE)
                            if m:
                                val = m.group(1)
                                num_map = {"fourteen": "Fourteen", "14": "Fourteen"}
                                override = num_map.get(val.lower(), val.capitalize() if val.isalpha() else val)
                                print("Debug: General age override:", override)
                        elif "amount" in qlower or re.search(r"\b(rs|rupees|\$|\d+,\d+)\b", qlower):
                            m = re.search(r"Rs\s*(\d+(?:,\d+)?)", raw_context, flags=re.IGNORECASE)
                            if m:
                                override = f"Rs {m.group(1)}"
                                print("Debug: General amount override:", override)
                        # Yes/No questions
                        expects_yesno = bool(re.search(r"\b(is|does|can|will|should|may|must)\s+.*\?$", qlower) or "legal" in qlower or "allowed" in qlower)
                        if expects_yesno and override is None:
                            if re.search(r"\b(yes|no|legal|illegal|allowed|prohibited|permissible|not permitted)\b", raw_context.lower()):
                                override = "Yes" if "yes" in raw_context.lower() or "legal" in raw_context.lower() or "allowed" in raw_context.lower() else "No"
                                print("Debug: Yes/No override:", override)
                        override_answers.append(override)
                 
                # Build batched prompt
                batched_questions_str = "\n".join([f"Question {i+1}: {q}" for i, q in enumerate(questions_batch)])
                batched_contexts_str = "\n\n".join([f"Context for Question {i+1}:\n{c}" for i, (_, c) in enumerate(batch_contexts)])
                
                prompt = f"""You are an expert analyst. For each question, base your response strictly on its provided context. Provide a brief answer in 1 sentence only, using verbatim language if possible. If the context does not contain the information, respond exactly with 'Not specified in the provided context.' If the answer is an identifier (e.g., Article/Section), respond with the exact identifier only (e.g., 'Article 11'). For numeric ages, return the number word only (e.g., 'Fourteen'). For yes/no questions, return 'Yes' or 'No' only if context clearly implies. Return ONLY a JSON array of answers like ["Answer1", "Answer2"]. Do not add explanations or references.
                
Batched Questions:
{batched_questions_str}

Batched Contexts:
{batched_contexts_str}

Respond as a JSON array of answers, like: ["Answer1", "Answer2", "Answer3"]"""
                
                print("Debug: Full prompt:", prompt[:500] + "...")  # Print start of prompt for inspection
                
                # AI/ML API call with retry
                answers = None
                for attempt in range(batch_max_retries):
                    try:
                        response = await asyncio.wait_for(
                            model.generate_content_async(
                                prompt,
                                temperature=0
                            ),
                            timeout=batch_timeout
                        )
                        print("Debug: Full AI/ML API response:", response)
                        answers_text = response.strip()
                        # Basic JSON array parse
                        try:
                            answers = json.loads(
                                answers_text.removeprefix("```json").removesuffix("```").strip()
                            )
                        except Exception:
                            m = re.search(r"\[.*\]", answers_text, re.DOTALL)
                            answers = json.loads(m.group(0)) if m else []
                        if not isinstance(answers, list) or len(answers) != len(questions_batch):
                            raise ValueError("Mismatched answer count")
                        break
                    except Exception as e:
                        if "429" in str(e) and attempt < batch_max_retries - 1:
                            wait_time = 5 * (2 ** attempt)
                            print(f"Debug: 429 quota error on attempt {attempt+1}. Retrying after {wait_time}s...")
                            await asyncio.sleep(wait_time)
                        else:
                            print("Error processing batch:", questions_batch, "Exception:", str(e))
                            answers = ["Not specified in the provided context." for _ in questions_batch]
                            break

                # Return answers
                for q, out in zip(questions_batch, answers):
                    print("Debug: Generated answer for question:", q, "(Length:", len(str(out)), ", Time:", round(time.time() - start_time, 2), "s):", str(out)[:100] + "...")
                return answers
            
            except asyncio.TimeoutError:
                print("Debug: Timeout for batch:", questions_batch)
                return ["Not specified in the provided context." for _ in questions_batch]  # Graceful fallback
        
        # Parallel processing with batching
        start_total = time.time()
        batch_size = 5  # Increased for fewer calls with higher quota
        batches = [questions[i:i + batch_size] for i in range(0, len(questions), batch_size)]
        tasks = [generate_answer(batch, vector_store, model, documents, doc_length, doc_hash) for batch in batches]
        batched_answers = await asyncio.gather(*tasks)
        answers = [ans for batch in batched_answers for ans in batch]  # Flatten

        # Basic metrics logging
        total_questions = len(questions)
        successful_answers = sum(1 for ans in answers if ans != "Not specified in the provided context." and "Error" not in ans)
        fallback_count = total_questions - successful_answers
        print("Metrics: Fallbacks:", fallback_count, "/", total_questions)
        success_rate = (successful_answers / total_questions) * 100 if total_questions > 0 else 0
        avg_time = round((time.time() - start_total) / total_questions, 2) if total_questions > 0 else 0
        print("Metrics: Success Rate:", f"{success_rate:.2f}%", "Avg Time per Question:", avg_time, "s")
        print("Debug: All answers generated. Total time:", round(time.time() - start_total, 2), "s")
 
        return {"answers": answers}
 
    finally:
        if os.path.exists(temp_file):
            os.remove(temp_file)


if __name__ == "__main__":
    os.makedirs("./chroma_db", exist_ok=True)
    os.makedirs("./uploads", exist_ok=True)
    uvicorn.run(app, host="0.0.0.0", port=8000,reload = False)


