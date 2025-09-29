"""
AWS-optimized RAG System for IntelliClaim
Replaces local ChromaDB with AWS RDS PostgreSQL + pgvector
"""

import asyncio
import json
import logging
import os
import time
from dataclasses import dataclass
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

from aws_database import vector_db, DocumentChunk
from aws_storage import s3_storage
from aws_config import get_aws_config
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import OpenAIEmbeddings
from langchain.schema import Document

logger = logging.getLogger(__name__)

@dataclass
class QueryContext:
    original_query: str
    structured_query: Dict[str, Any] = None
    retrieved_docs: List[DocumentChunk] = None
    reasoning_chain: List[str] = None
    decision: Dict[str, Any] = None
    confidence_score: float = 0.0
    explanation: Dict[str, Any] = None

@dataclass
class DecisionResponse:
    decision: str
    amount: Optional[float] = None
    justification: str = ""
    clause_mappings: List[Dict[str, str]] = None
    confidence_score: float = 0.0
    processing_time: float = 0.0
    audit_trail: List[str] = None

class AWSRAGSystem:
    """AWS-optimized RAG system using RDS + pgvector"""
    
    def __init__(self):
        self.aws_config = get_aws_config()
        self.embedding_manager = OpenAIEmbeddings(
            openai_api_key=os.getenv("OPENAI_API_KEY"),
            model="text-embedding-3-small",
            chunk_size=1000
        )
        self.doc_processor = RecursiveCharacterTextSplitter(
            chunk_size=2000,
            chunk_overlap=100
        )
        self.initialized = False
    
    async def initialize(self):
        """Initialize the RAG system with AWS services"""
        try:
            await vector_db.initialize()
            self.initialized = True
            logger.info("AWS RAG system initialized successfully")
        except Exception as e:
            logger.error(f"Failed to initialize AWS RAG system: {e}")
            raise
    
    async def add_document(self, file_path: str, file_name: str = None) -> Dict[str, Any]:
        """Add document to AWS storage and vector database"""
        if not self.initialized:
            await self.initialize()
        
        try:
            # Process document locally
            documents = await self._process_document(file_path)
            
            if not documents:
                return {"status": "error", "message": "No text could be extracted from the document"}
            
            # Upload to S3
            s3_result = s3_storage.upload_document(file_path, file_name)
            if s3_result["status"] != "success":
                return s3_result
            
            # Convert to DocumentChunk objects
            chunks = []
            for i, doc in enumerate(documents):
                # Generate embedding
                embedding = await self._generate_embedding(doc.page_content)
                
                chunk = DocumentChunk(
                    id=f"{file_name}_{i}",
                    content=doc.page_content,
                    metadata=doc.metadata,
                    embedding=embedding,
                    source=file_name or os.path.basename(file_path),
                    chunk_id=str(i)
                )
                chunks.append(chunk)
            
            # Add to AWS database
            result = await vector_db.add_document(file_path, file_name or os.path.basename(file_path), chunks)
            
            # Log query for analytics
            await vector_db.log_query(
                query="Document upload",
                structured_query={"action": "upload", "file": file_name},
                retrieved_chunks=0,
                decision="UPLOADED",
                confidence_score=100.0,
                processing_time=0.0
            )
            
            return result
            
        except Exception as e:
            logger.error(f"Failed to add document {file_path}: {e}")
            return {"status": "error", "message": str(e)}
    
    async def process_query(self, query: str) -> DecisionResponse:
        """Process query using AWS RAG system"""
        if not self.initialized:
            await self.initialize()
        
        start_time = datetime.now()
        
        try:
            # Create query context
            context = QueryContext(original_query=query)
            
            # Step 1: Query understanding
            context = await self._understand_query(context)
            
            # Step 2: Retrieve relevant documents
            context = await self._retrieve_documents(context)
            
            # Step 3: Make decision
            context = await self._make_decision(context)
            
            # Step 4: Generate explanation
            context = await self._generate_explanation(context)
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            # Log query for analytics
            await vector_db.log_query(
                query=query,
                structured_query=context.structured_query or {},
                retrieved_chunks=len(context.retrieved_docs or []),
                decision=context.decision.get('decision', 'UNKNOWN') if context.decision else 'UNKNOWN',
                confidence_score=context.decision.get('confidence_score', 0) if context.decision else 0,
                processing_time=processing_time
            )
            
            return DecisionResponse(
                decision=context.decision.get('decision', 'ERROR') if context.decision else 'ERROR',
                amount=context.decision.get('amount') if context.decision else None,
                justification=context.decision.get('justification', 'An unknown error occurred.') if context.decision else 'An unknown error occurred.',
                confidence_score=context.decision.get('confidence_score', 0) if context.decision else 0,
                clause_mappings=context.explanation.get('clause_mappings', []) if context.explanation else [],
                audit_trail=context.explanation.get('audit_trail', []) if context.explanation else [],
                processing_time=round(processing_time, 2)
            )
            
        except Exception as e:
            logger.error(f"Error in process_query: {e}")
            processing_time = (datetime.now() - start_time).total_seconds()
            return DecisionResponse(
                decision="ERROR",
                justification=f"Processing failed: {str(e)}",
                confidence_score=0,
                clause_mappings=[],
                audit_trail=["Error occurred during processing"],
                processing_time=round(processing_time, 2)
            )
    
    async def _process_document(self, file_path: str) -> List[Document]:
        """Process document and extract text"""
        try:
            import PyPDF2
            from docx import Document as DocxDocument
            from email import policy as email_policy
            from email.parser import BytesParser as EmailBytesParser
            
            filename = os.path.basename(file_path)
            ext = os.path.splitext(filename)[1].lower()
            
            raw_text = ""
            
            if ext == '.pdf':
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    parts = []
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text and text.strip():
                            parts.append(text)
                    raw_text = "\n\n".join(parts)
            
            elif ext == '.docx':
                doc = DocxDocument(file_path)
                raw_text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            
            elif ext == '.eml':
                with open(file_path, 'rb') as fp:
                    msg = EmailBytesParser(policy=email_policy.default).parse(fp)
                parts = [
                    msg.get('subject', ''),
                    msg.get_body(preferencelist=('plain', 'html')).get_content() if msg.get_body() else ''
                ]
                raw_text = "\n\n".join(p for p in parts if p)
            
            else:
                logger.warning(f"Unsupported file type: {ext}")
                return []
            
            if not raw_text.strip():
                return []
            
            # Split into chunks
            chunks = self.doc_processor.split_text(raw_text)
            
            # Create Document objects
            documents = []
            for i, chunk in enumerate(chunks):
                documents.append(Document(
                    page_content=chunk,
                    metadata={
                        "source": filename,
                        "chunk_id": str(i),
                        "content_type": "text"
                    }
                ))
            
            return documents
            
        except Exception as e:
            logger.error(f"Document processing failed for {file_path}: {e}")
            return []
    
    async def _generate_embedding(self, text: str) -> List[float]:
        """Generate embedding for text"""
        try:
            # Use OpenAI embeddings
            embedding = await asyncio.to_thread(self.embedding_manager.embed_query, text)
            return embedding
        except Exception as e:
            logger.error(f"Failed to generate embedding: {e}")
            # Return dummy embedding if OpenAI fails
            return [0.0] * 1536
    
    async def _understand_query(self, context: QueryContext) -> QueryContext:
        """Extract structured information from query"""
        try:
            # Simple entity extraction (you can enhance this with LLM)
            query_lower = context.original_query.lower()
            
            # Extract age
            import re
            age_match = re.search(r'(\d+)[\s-]*(?:year|yr)s?[\s-]*old', query_lower)
            age = int(age_match.group(1)) if age_match else None
            
            # Extract gender
            gender = None
            if 'male' in query_lower or 'm' in query_lower:
                gender = 'male'
            elif 'female' in query_lower or 'woman' in query_lower:
                gender = 'female'
            
            # Extract location
            locations = ['pune', 'mumbai', 'delhi', 'chennai', 'bangalore', 'hyderabad', 'kolkata']
            location = None
            for loc in locations:
                if loc in query_lower:
                    location = loc.title()
                    break
            
            # Extract procedure
            procedures = ['surgery', 'treatment', 'operation', 'dental', 'cataract', 'heart', 'knee']
            procedure = None
            for proc in procedures:
                if proc in query_lower:
                    procedure = proc
                    break
            
            context.structured_query = {
                "age": age,
                "gender": gender,
                "location": location,
                "procedure": procedure,
                "policy_duration_months": 12,  # Default
                "intent": "claim_eligibility"
            }
            
        except Exception as e:
            logger.error(f"Query understanding failed: {e}")
            context.structured_query = {
                "age": None,
                "gender": None,
                "location": None,
                "procedure": None,
                "policy_duration_months": 12,
                "intent": "claim_eligibility"
            }
        
        return context
    
    async def _retrieve_documents(self, context: QueryContext) -> QueryContext:
        """Retrieve relevant documents using vector search"""
        try:
            # Create search query
            search_query = " ".join(filter(None, [
                context.structured_query.get('procedure'),
                context.structured_query.get('location'),
                "insurance policy",
                "coverage"
            ]))
            
            if not search_query.strip():
                search_query = context.original_query
            
            # Generate embedding for search
            query_embedding = await self._generate_embedding(search_query)
            
            # Perform vector search
            results = await vector_db.similarity_search(query_embedding, k=5)
            
            context.retrieved_docs = results
            logger.info(f"Retrieved {len(results)} documents")
            
        except Exception as e:
            logger.error(f"Document retrieval failed: {e}")
            context.retrieved_docs = []
        
        return context
    
    async def _make_decision(self, context: QueryContext) -> QueryContext:
        """Make decision based on retrieved documents"""
        try:
            if not context.retrieved_docs:
                context.decision = {
                    "decision": "PENDING",
                    "justification": "No relevant policy documents found. Please upload policy documents for assessment.",
                    "amount": None,
                    "confidence_score": 20
                }
                return context
            
            # Simple rule-based decision making
            # You can enhance this with LLM-based decision making
            age = context.structured_query.get('age')
            procedure = context.structured_query.get('procedure')
            
            decision = "PENDING"
            justification = "Requires manual review"
            amount = None
            confidence = 50
            
            if age and age > 65:
                if procedure in ['cataract', 'heart']:
                    decision = "APPROVED"
                    justification = f"Standard procedure for age {age} approved under policy"
                    amount = 50000  # Example amount
                    confidence = 80
                else:
                    decision = "PENDING"
                    justification = f"Age {age} requires additional medical review for {procedure}"
                    confidence = 60
            elif age and age < 30:
                decision = "APPROVED"
                justification = f"Young age {age} with low risk for {procedure}"
                amount = 30000  # Example amount
                confidence = 75
            else:
                decision = "PENDING"
                justification = f"Standard review required for age {age} and {procedure}"
                confidence = 55
            
            context.decision = {
                "decision": decision,
                "justification": justification,
                "amount": amount,
                "confidence_score": confidence
            }
            
        except Exception as e:
            logger.error(f"Decision making failed: {e}")
            context.decision = {
                "decision": "ERROR",
                "justification": f"Decision processing failed: {str(e)}",
                "amount": None,
                "confidence_score": 0
            }
        
        return context
    
    async def _generate_explanation(self, context: QueryContext) -> QueryContext:
        """Generate explanation and audit trail"""
        try:
            clause_mappings = []
            if context.retrieved_docs:
                for doc in context.retrieved_docs[:3]:  # Top 3 documents
                    clause_mappings.append({
                        "clause_text": doc.content[:200] + "...",
                        "source": doc.source
                    })
            
            audit_trail = [
                "Extracted key entities from query",
                f"Retrieved {len(context.retrieved_docs or [])} relevant policy documents",
                f"Applied rules to make a {context.decision.get('decision', 'unknown')} decision" if context.decision else "Decision processing failed"
            ]
            
            context.explanation = {
                "clause_mappings": clause_mappings,
                "audit_trail": audit_trail
            }
            
        except Exception as e:
            logger.error(f"Explanation generation failed: {e}")
            context.explanation = {
                "clause_mappings": [],
                "audit_trail": ["Error occurred during explanation generation"]
            }
        
        return context
    
    async def get_system_stats(self) -> Dict[str, Any]:
        """Get system statistics"""
        try:
            return await vector_db.get_system_stats()
        except Exception as e:
            logger.error(f"Failed to get system stats: {e}")
            return {
                "total_documents": 0,
                "total_chunks": 0,
                "system_status": "error",
                "database_status": "error",
                "error": str(e)
            }
    
    async def list_documents(self) -> List[Dict[str, Any]]:
        """List all documents in the system"""
        try:
            return await vector_db.list_documents()
        except Exception as e:
            logger.error(f"Failed to list documents: {e}")
            return []
    
    async def delete_document(self, file_name: str) -> Dict[str, Any]:
        """Delete document from the system"""
        try:
            return await vector_db.delete_document(file_name)
        except Exception as e:
            logger.error(f"Failed to delete document {file_name}: {e}")
            return {"status": "error", "message": str(e)}

# Global RAG system instance
aws_rag_system = AWSRAGSystem()
